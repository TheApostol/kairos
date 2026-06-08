"""Convert MANUAL_KAIROS.md to DOCX and PDF."""

import re
from pathlib import Path

MD_PATH = Path(__file__).parent.parent / "MANUAL_KAIROS.md"
OUT_DIR = Path(__file__).parent.parent
DOCX_PATH = OUT_DIR / "MANUAL_KAIROS.docx"
PDF_PATH = OUT_DIR / "MANUAL_KAIROS.pdf"

# ──────────────────────────────────────────────────────────
# Colours / brand
# ──────────────────────────────────────────────────────────
GREEN = "2D6A4F"    # Kairos dark green
GOLD  = "C9A040"    # Kairos gold accent
DARK  = "1A1A2E"    # near-black for headings
GRAY  = "64748B"    # body text
LGRAY = "F1F5F9"    # table row shading

# ──────────────────────────────────────────────────────────
# 1.  DOCX
# ──────────────────────────────────────────────────────────
def build_docx(lines):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Helper: set paragraph font
    def _font(run, size, bold=False, color=None, italic=False):
        run.font.name  = "Calibri"
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.italic = italic
        if color:
            r, g, b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
            run.font.color.rgb = RGBColor(r, g, b)

    def _shade_cell(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _set_cell_border(cell, **kwargs):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top","left","bottom","right","insideH","insideV"):
            if side in kwargs:
                tag = OxmlElement(f"w:{side}")
                tag.set(qn("w:val"),  kwargs[side].get("val","single"))
                tag.set(qn("w:sz"),   kwargs[side].get("sz","4"))
                tag.set(qn("w:color"),kwargs[side].get("color","auto"))
                tcBorders.append(tag)
        tcPr.append(tcBorders)

    # ── Cover page ──────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(80)
    r = cover.add_run("KAIROS CRM")
    _font(r, 36, bold=True, color=GREEN)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Manual de Usuario")
    _font(r2, 18, color=GOLD)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = sub2.add_run("Kairos Distribuidora · Guía completa del sistema")
    _font(r3, 11, color=GRAY)

    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = sub3.add_run("Junio 2026  ·  v2.0")
    _font(r4, 10, italic=True, color=GRAY)

    doc.add_page_break()

    # ── Body ──────────────────────────────────────────
    i = 0
    in_table   = False
    table_rows = []
    in_code    = False
    code_lines = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        # strip separator rows (---|---)
        data_rows = [r for r in table_rows if not re.match(r"^\|[\s\-|:]+\|$", r)]
        if not data_rows:
            table_rows = []
            return
        parsed = []
        for row in data_rows:
            cells = [c.strip() for c in row.strip("|").split("|")]
            parsed.append(cells)
        cols  = max(len(r) for r in parsed)
        tbl   = doc.add_table(rows=len(parsed), cols=cols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for ri, row in enumerate(parsed):
            for ci, cell_text in enumerate(row):
                if ci >= cols:
                    break
                cell = tbl.cell(ri, ci)
                p    = cell.paragraphs[0]
                p.clear()
                is_header = ri == 0
                if is_header:
                    _shade_cell(cell, GREEN)
                elif ri % 2 == 0:
                    _shade_cell(cell, LGRAY)
                run = p.add_run(cell_text)
                _font(run, 9, bold=is_header,
                      color="FFFFFF" if is_header else DARK)
        doc.add_paragraph()
        table_rows = []

    def flush_code():
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        # light grey background via shading on paragraph — use a 1-cell table
        tbl  = doc.add_table(rows=1, cols=1)
        cell = tbl.cell(0,0)
        _shade_cell(cell, "F8FAFC")
        p2   = cell.paragraphs[0]
        p2.clear()
        for ln in code_lines:
            run = p2.add_run(ln + "\n")
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
            r_el = run._r
            rPr  = r_el.get_or_add_rPr()
        doc.add_paragraph()
        code_lines = []

    while i < len(lines):
        line = lines[i]

        # ── fenced code block ──────────────────────────
        if line.startswith("```"):
            if not in_code:
                in_code = True
                i += 1
                continue
            else:
                in_code = False
                flush_code()
                i += 1
                continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── table ──────────────────────────────────────
        if line.startswith("|"):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        if in_table:
            flush_table()
            in_table = False

        # ── horizontal rule ────────────────────────────
        if re.match(r"^---+$", line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),  "single")
            bot.set(qn("w:sz"),   "4")
            bot.set(qn("w:color"), GREEN)
            pBdr.append(bot)
            pPr.append(pBdr)
            i += 1
            continue

        # ── headings ───────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = m.group(2)
            # strip anchor links  [text](#anchor) → text
            text  = re.sub(r"\[([^\]]+)\]\(#[^)]+\)", r"\1", text)
            p = doc.add_paragraph()
            if level == 1:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(6)
                r = p.add_run(text.upper())
                _font(r, 20, bold=True, color=GREEN)
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bot  = OxmlElement("w:bottom")
                bot.set(qn("w:val"),  "single")
                bot.set(qn("w:sz"),   "8")
                bot.set(qn("w:color"), GREEN)
                pBdr.append(bot)
                pPr.append(pBdr)
            elif level == 2:
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after  = Pt(4)
                r = p.add_run(text)
                _font(r, 15, bold=True, color=GREEN)
            elif level == 3:
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after  = Pt(2)
                r = p.add_run(text)
                _font(r, 12, bold=True, color=DARK)
            else:
                p.paragraph_format.space_before = Pt(6)
                r = p.add_run(text)
                _font(r, 11, bold=True, color=GOLD)
            i += 1
            continue

        # ── bullet list ────────────────────────────────
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            text   = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent  = Cm(0.5 + indent * 0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            _add_inline(p, text, 10)
            i += 1
            continue

        # ── numbered list ──────────────────────────────
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            text   = m.group(2)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent  = Cm(0.5 + indent * 0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            _add_inline(p, text, 10)
            i += 1
            continue

        # ── blank line ─────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── normal paragraph ───────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        _add_inline(p, line, 10)
        i += 1

    flush_table()
    flush_code()

    doc.save(DOCX_PATH)
    print(f"✓ DOCX → {DOCX_PATH}")


def _add_inline(paragraph, text, size):
    """Add text with **bold**, `code`, and *italic* inline markers."""
    from docx.shared import Pt, RGBColor

    # Remove markdown link syntax [text](url) → text
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)

    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.font.bold = True
            r.font.name = "Calibri"
            r.font.size = Pt(size)
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(size - 0.5)
            r.font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F)
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.font.italic = True
            r.font.name   = "Calibri"
            r.font.size   = Pt(size)
        else:
            r = paragraph.add_run(part)
            r.font.name = "Calibri"
            r.font.size = Pt(size)


# ──────────────────────────────────────────────────────────
# 2.  PDF  (ReportLab)
# ──────────────────────────────────────────────────────────
def build_pdf(lines):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm, mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable, Preformatted
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    W, H = A4

    # Colours
    cGreen = colors.HexColor("#2D6A4F")
    cGold  = colors.HexColor("#C9A040")
    cDark  = colors.HexColor("#1A1A2E")
    cGray  = colors.HexColor("#64748B")
    cLGray = colors.HexColor("#F1F5F9")
    cWhite = colors.white

    styles = getSampleStyleSheet()

    def S(name, **kw):
        return ParagraphStyle(name, **kw)

    s_h1 = S("H1", fontSize=20, textColor=cGreen, fontName="Helvetica-Bold",
              spaceBefore=16, spaceAfter=6, leading=24)
    s_h2 = S("H2", fontSize=14, textColor=cGreen, fontName="Helvetica-Bold",
              spaceBefore=14, spaceAfter=4, leading=18)
    s_h3 = S("H3", fontSize=11, textColor=cDark,  fontName="Helvetica-Bold",
              spaceBefore=10, spaceAfter=3, leading=14)
    s_h4 = S("H4", fontSize=10, textColor=cGold,  fontName="Helvetica-Bold",
              spaceBefore=6,  spaceAfter=2, leading=13)
    s_body = S("Body", fontSize=9.5, textColor=cDark, fontName="Helvetica",
               spaceBefore=2, spaceAfter=4, leading=13)
    s_bullet = S("Bullet", fontSize=9.5, textColor=cDark, fontName="Helvetica",
                 leftIndent=14, firstLineIndent=-10,
                 spaceBefore=1, spaceAfter=1, leading=13)
    s_code  = S("Code",  fontSize=8,   textColor=cDark, fontName="Courier",
                backColor=colors.HexColor("#F8FAFC"),
                leftIndent=12, rightIndent=12,
                spaceBefore=4, spaceAfter=4, leading=12)
    s_cover_title = S("CT", fontSize=36, textColor=cGreen, fontName="Helvetica-Bold",
                      alignment=TA_CENTER, spaceBefore=100, spaceAfter=8)
    s_cover_sub   = S("CS", fontSize=16, textColor=cGold,  fontName="Helvetica-Bold",
                      alignment=TA_CENTER, spaceAfter=6)
    s_cover_body  = S("CB", fontSize=11, textColor=cGray,  fontName="Helvetica",
                      alignment=TA_CENTER, spaceAfter=4)

    def md_to_para(text, style):
        """Convert minimal markdown inline (bold, code, italic) to ReportLab XML."""
        # links
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        # escape XML chars first (except our tags)
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # **bold**
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        # `code`
        text = re.sub(r"`(.+?)`", r'<font name="Courier" color="#2D6A4F">\1</font>', text)
        # *italic*
        text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
        return Paragraph(text, style)

    story = []

    # Cover
    story.append(Paragraph("KAIROS CRM", s_cover_title))
    story.append(Paragraph("Manual de Usuario", s_cover_sub))
    story.append(Paragraph("Kairos Distribuidora · Guía completa del sistema", s_cover_body))
    story.append(Paragraph("Junio 2026  ·  v2.0", s_cover_body))
    story.append(PageBreak())

    in_table   = False
    table_rows = []
    in_code    = False
    code_lines = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        data_rows = [r for r in table_rows if not re.match(r"^\|[\s\-|:]+\|$", r)]
        if not data_rows:
            table_rows.clear()
            return
        parsed = []
        for row in data_rows:
            cells = [c.strip() for c in row.strip("|").split("|")]
            parsed.append(cells)
        cols = max(len(r) for r in parsed)
        # normalise row widths
        for row in parsed:
            while len(row) < cols:
                row.append("")
        col_w = (W - 5*cm) / cols

        def cell_para(txt, is_header):
            txt = txt.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            txt = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", txt)
            txt = re.sub(r"`(.+?)`", r'<font name="Courier">\1</font>', txt)
            s = ParagraphStyle("tc", fontSize=8, textColor=cWhite if is_header else cDark,
                               fontName="Helvetica-Bold" if is_header else "Helvetica",
                               leading=11)
            return Paragraph(txt, s)

        tdata = [[cell_para(c, ri==0) for c in row] for ri, row in enumerate(parsed)]
        ts = TableStyle([
            ("BACKGROUND", (0,0), (-1,0), cGreen),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [cWhite, cLGray]),
            ("GRID",       (0,0), (-1,-1), 0.25, colors.HexColor("#CBD5E1")),
            ("TOPPADDING", (0,0), (-1,-1), 4),
            ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ("LEFTPADDING",   (0,0), (-1,-1), 6),
            ("RIGHTPADDING",  (0,0), (-1,-1), 6),
            ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
        ])
        t = Table(tdata, colWidths=[col_w]*cols, style=ts, repeatRows=1)
        story.append(t)
        story.append(Spacer(1, 6))
        table_rows.clear()

    def flush_code():
        nonlocal code_lines
        if not code_lines:
            return
        raw = "\n".join(code_lines)
        story.append(Preformatted(raw, s_code))
        story.append(Spacer(1, 4))
        code_lines.clear()

    for line in lines:
        # code fence
        if line.startswith("```"):
            if not in_code:
                in_code = True
            else:
                in_code = False
                flush_code()
            continue
        if in_code:
            code_lines.append(line)
            continue

        # table
        if line.startswith("|"):
            in_table = True
            table_rows.append(line)
            continue
        if in_table:
            flush_table()
            in_table = False

        # HR
        if re.match(r"^---+$", line.strip()):
            story.append(HRFlowable(width="100%", thickness=1, color=cGreen,
                                    spaceAfter=6, spaceBefore=4))
            continue

        # headings
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = m.group(2)
            text  = re.sub(r"\[([^\]]+)\]\(#[^)]+\)", r"\1", text)
            st    = [s_h1, s_h2, s_h3, s_h4][level-1]
            if level == 1:
                story.append(HRFlowable(width="100%", thickness=2, color=cGreen,
                                        spaceAfter=2, spaceBefore=14))
            story.append(Paragraph(text, st))
            if level == 1:
                story.append(HRFlowable(width="100%", thickness=2, color=cGreen,
                                        spaceAfter=4, spaceBefore=2))
            continue

        # bullet
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            text   = m.group(2)
            s = ParagraphStyle("bl", parent=s_bullet, leftIndent=14 + indent*12,
                               bulletIndent=4 + indent*12)
            story.append(md_to_para("• " + text, s))
            continue

        # numbered
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            num    = re.match(r"^(\s*)(\d+)\.", line).group(2)
            text   = m.group(2)
            s = ParagraphStyle("nl", parent=s_bullet, leftIndent=14 + indent*12,
                               bulletIndent=4 + indent*12)
            story.append(md_to_para(f"{num}. {text}", s))
            continue

        # blank
        if line.strip() == "":
            story.append(Spacer(1, 3))
            continue

        # normal
        story.append(md_to_para(line, s_body))

    flush_table()
    flush_code()

    def _header_footer(canvas, doc):
        canvas.saveState()
        # Header bar
        canvas.setFillColor(cGreen)
        canvas.rect(0, H - 1.2*cm, W, 1.2*cm, fill=1, stroke=0)
        canvas.setFillColor(cWhite)
        canvas.setFont("Helvetica-Bold", 9)
        canvas.drawString(2*cm, H - 0.8*cm, "KAIROS CRM  ·  Manual de Usuario")
        # Footer
        canvas.setFillColor(cGray)
        canvas.setFont("Helvetica", 8)
        canvas.drawString(2*cm, 0.8*cm, "Kairos Distribuidora  ·  Junio 2026  ·  v2.0")
        canvas.drawRightString(W - 2*cm, 0.8*cm, f"Página {doc.page}")
        canvas.restoreState()

    pdf = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        topMargin=1.8*cm, bottomMargin=1.8*cm,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        title="Manual Kairos CRM",
        author="Kairos Distribuidora",
    )
    pdf.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
    print(f"✓ PDF  → {PDF_PATH}")


# ──────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────
if __name__ == "__main__":
    raw   = MD_PATH.read_text(encoding="utf-8")
    lines = raw.splitlines()

    # Skip the very first H1 line (used as cover) in the body
    body_lines = [l for l in lines if not l.startswith("# Manual de Usuario")]

    build_docx(body_lines)
    build_pdf(body_lines)
    print("Done.")
