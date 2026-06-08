"""Generates GUIA_KAIROS_CRM.docx — a fully visual Word guide."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colours ──────────────────────────────────────────────────────────
BROWN_DARK   = RGBColor(0x4A, 0x37, 0x28)   # #4A3728
BROWN_MID    = RGBColor(0x6B, 0x4F, 0x3A)   # #6B4F3A
GOLD         = RGBColor(0xC9, 0xA0, 0x40)   # #C9A040
GOLD_LIGHT   = RGBColor(0xFA, 0xF7, 0xF2)   # #FAF7F2
SIDEBAR_BG   = RGBColor(0x2C, 0x1F, 0x16)   # #2C1F16
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
SLATE_LIGHT  = RGBColor(0xF8, 0xF9, 0xFA)
SLATE_BORDER = RGBColor(0xE2, 0xE8, 0xF0)
GREEN        = RGBColor(0x16, 0xA3, 0x4A)
BLUE         = RGBColor(0x25, 0x63, 0xEB)
RED          = RGBColor(0xDC, 0x26, 0x26)
AMBER        = RGBColor(0xD9, 0x77, 0x06)


# ── Low-level XML helpers ──────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    """Fill a table cell with a solid background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = str(rgb)   # RGBColor.__str__ returns 'RRGGBB'
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val', 'single'))
            el.set(qn('w:sz'),    str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def no_borders(table):
    """Remove all borders from a table."""
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(int(width_cm * 567)))  # 567 twips/cm
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)


def add_run(para, text, bold=False, italic=False, color=None, size=None, font=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if font:
        run.font.name = font
    return run


def set_para_spacing(para, before=0, after=0, line=None):
    pPr  = para._p.get_or_add_pPr()
    spg  = OxmlElement('w:spacing')
    spg.set(qn('w:before'), str(int(before * 20)))
    spg.set(qn('w:after'),  str(int(after  * 20)))
    if line:
        spg.set(qn('w:line'),     str(int(line * 240)))
        spg.set(qn('w:lineRule'), 'auto')
    pPr.append(spg)


def set_left_indent(para, cm):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(int(cm * 567)))
    pPr.append(ind)


def add_page_break(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(docx_break_type())


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml    import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br   # we won't use this helper as-is; see below


def page_break(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run  = para.add_run()
    br   = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


# ── Reusable section builders ──────────────────────────────────────────────

def section_header(doc, number, title, subtitle=''):
    """Gold accent bar + section title."""
    # Accent stripe via a 1-row, 2-col table
    t = doc.add_table(rows=1, cols=2)
    no_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Gold bar (narrow)
    bar = t.cell(0, 0)
    set_cell_bg(bar, GOLD)
    bar_w = bar._tc.get_or_add_tcPr()
    w_el  = OxmlElement('w:tcW')
    w_el.set(qn('w:w'),    '280')
    w_el.set(qn('w:type'), 'dxa')
    bar_w.append(w_el)
    bar.paragraphs[0].paragraph_format.space_before = Pt(2)
    bar.paragraphs[0].paragraph_format.space_after  = Pt(2)

    # Title cell
    cell = t.cell(0, 1)
    set_cell_bg(cell, BROWN_DARK)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run(f'{number}.  ')
    r1.font.color.rgb = GOLD
    r1.font.size      = Pt(13)
    r1.bold           = True
    r2 = p.add_run(title.upper())
    r2.font.color.rgb = WHITE
    r2.font.size      = Pt(13)
    r2.bold           = True

    if subtitle:
        sp = cell.add_paragraph(subtitle)
        sp.paragraph_format.left_indent  = Cm(0.3)
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(6)
        sp.runs[0].font.color.rgb = GOLD_LIGHT
        sp.runs[0].font.size      = Pt(9)
        sp.runs[0].italic         = True

    doc.add_paragraph()   # breathing room


def h2(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=10, after=3)
    r = p.add_run(text)
    r.bold            = True
    r.font.size       = Pt(11)
    r.font.color.rgb  = BROWN_DARK
    # Gold underline via bottom border on the paragraph
    pPr      = p._p.get_or_add_pPr()
    pBdr     = OxmlElement('w:pBdr')
    bottom   = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C9A040')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def body(doc, text, indent=False):
    p = doc.add_paragraph()
    set_para_spacing(p, before=2, after=2, line=1.15)
    if indent:
        set_left_indent(p, 0.6)
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if level:
        p.paragraph_format.left_indent = Cm(1.2 * level)
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p


def tip_box(doc, text, color=GOLD, bg=None):
    """Highlighted callout box."""
    if bg is None:
        bg = RGBColor(
            min(255, color[0] + 220),
            min(255, color[1] + 200),
            min(255, color[2] + 180),
        )
    t = doc.add_table(rows=1, cols=2)
    no_borders(t)

    stripe = t.cell(0, 0)
    set_cell_bg(stripe, color)
    sw = stripe._tc.get_or_add_tcPr()
    we = OxmlElement('w:tcW')
    we.set(qn('w:w'), '160')
    we.set(qn('w:type'), 'dxa')
    sw.append(we)
    stripe.paragraphs[0].add_run(' ')

    content = t.cell(0, 1)
    set_cell_bg(content, bg)
    p = content.paragraphs[0]
    p.paragraph_format.left_indent  = Cm(0.25)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run(text)
    r.font.size      = Pt(9.5)
    r.font.italic    = True
    r.font.color.rgb = BROWN_DARK

    doc.add_paragraph()


def data_table(doc, headers, rows, col_widths=None):
    """Styled data table with gold header row."""
    ncols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=ncols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = t.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, BROWN_DARK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(hdr)
        r.bold            = True
        r.font.color.rgb  = WHITE
        r.font.size       = Pt(9.5)

    # Data rows
    for ri, row_data in enumerate(rows):
        tr = t.rows[ri + 1]
        bg = GOLD_LIGHT if ri % 2 == 0 else WHITE
        for ci, cell_text in enumerate(row_data):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            if isinstance(cell_text, tuple):
                text, bold, color = cell_text
                r = p.add_run(text)
                r.bold            = bold
                r.font.size       = Pt(9.5)
                r.font.color.rgb  = color
            else:
                r = p.add_run(str(cell_text))
                r.font.size = Pt(9.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            set_col_width(t, i, w)

    doc.add_paragraph()
    return t


def flow_table(doc, steps):
    """Horizontal flow diagram using table cells with arrows."""
    n = len(steps)
    # Each step = 1 col, arrows between = 1 col
    ncols = n * 2 - 1
    t = doc.add_table(rows=1, cols=ncols)
    no_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (icon, label) in enumerate(steps):
        col = i * 2
        cell = t.cell(0, col)
        set_cell_bg(cell, BROWN_DARK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(3)
        r1 = p.add_run(icon + '\n')
        r1.font.size      = Pt(14)
        r1.font.color.rgb = GOLD
        r1.bold = True
        r2 = p.add_run(label)
        r2.font.size      = Pt(8)
        r2.font.color.rgb = WHITE
        r2.bold = True

        # set width
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        we = OxmlElement('w:tcW')
        we.set(qn('w:w'), '1300')
        we.set(qn('w:type'), 'dxa')
        tcPr.append(we)

        if i < n - 1:
            arrow_cell = t.cell(0, col + 1)
            set_cell_bg(arrow_cell, WHITE)
            ap = arrow_cell.paragraphs[0]
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ap.paragraph_format.space_before = Pt(8)
            ar = ap.add_run('▶')
            ar.font.size      = Pt(14)
            ar.font.color.rgb = GOLD
            # narrow width
            tc2 = arrow_cell._tc
            tcPr2 = tc2.get_or_add_tcPr()
            we2 = OxmlElement('w:tcW')
            we2.set(qn('w:w'), '400')
            we2.set(qn('w:type'), 'dxa')
            tcPr2.append(we2)

    doc.add_paragraph()


def status_badge_table(doc, statuses):
    """A row of coloured status badges."""
    t = doc.add_table(rows=1, cols=len(statuses))
    no_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    colors_map = {
        'green':  (RGBColor(0xDC, 0xFC, 0xE7), GREEN),
        'blue':   (RGBColor(0xDB, 0xEA, 0xFE), BLUE),
        'amber':  (RGBColor(0xFF, 0xF7, 0xED), AMBER),
        'red':    (RGBColor(0xFE, 0xE2, 0xE2), RED),
        'gray':   (RGBColor(0xF1, 0xF5, 0xF9), RGBColor(0x64, 0x74, 0x8B)),
    }
    for i, (label, color_key) in enumerate(statuses):
        bg, fg = colors_map.get(color_key, (GOLD_LIGHT, BROWN_DARK))
        cell = t.cell(0, i)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(label)
        r.font.color.rgb = fg
        r.font.size      = Pt(9)
        r.bold           = True
    doc.add_paragraph()


def wizard_step_box(doc, step_num, step_title, description):
    """Visual wizard step box."""
    t = doc.add_table(rows=1, cols=2)
    no_borders(t)

    # Step number circle (using cell bg)
    num_cell = t.cell(0, 0)
    set_cell_bg(num_cell, GOLD)
    tc = num_cell._tc
    tcPr = tc.get_or_add_tcPr()
    we = OxmlElement('w:tcW')
    we.set(qn('w:w'), '600')
    we.set(qn('w:type'), 'dxa')
    tcPr.append(we)
    p = num_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(str(step_num))
    r.font.size      = Pt(16)
    r.font.color.rgb = WHITE
    r.bold           = True

    # Step content
    content_cell = t.cell(0, 1)
    set_cell_bg(content_cell, GOLD_LIGHT)
    p2 = content_cell.paragraphs[0]
    p2.paragraph_format.left_indent  = Cm(0.3)
    p2.paragraph_format.space_before = Pt(6)
    p2.paragraph_format.space_after  = Pt(2)
    r2 = p2.add_run(step_title)
    r2.font.size      = Pt(11)
    r2.font.color.rgb = BROWN_DARK
    r2.bold           = True

    p3 = content_cell.add_paragraph(description)
    p3.paragraph_format.left_indent  = Cm(0.3)
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(6)
    p3.runs[0].font.size      = Pt(9.5)
    p3.runs[0].font.color.rgb = BROWN_MID

    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════
#  DOCUMENT
# ══════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
from docx.oxml import OxmlElement
section = doc.sections[0]
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# Default font
from docx.oxml.ns import qn
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)


# ── COVER PAGE ─────────────────────────────────────────────────────────────

# Full-width coloured banner via table
cover = doc.add_table(rows=1, cols=1)
no_borders(cover)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cc = cover.cell(0, 0)
set_cell_bg(cc, SIDEBAR_BG)
# width
tc = cc._tc
tcPr = tc.get_or_add_tcPr()
we = OxmlElement('w:tcW')
we.set(qn('w:w'), '10080')   # ~17.8 cm
we.set(qn('w:type'), 'dxa')
tcPr.append(we)

p_logo = cc.paragraphs[0]
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_logo.paragraph_format.space_before = Pt(40)
p_logo.paragraph_format.space_after  = Pt(6)
r = p_logo.add_run('KAIROS')
r.font.size       = Pt(42)
r.font.color.rgb  = GOLD
r.bold            = True

p_sub = cc.add_paragraph('Distribuidora')
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(0)
r2 = p_sub.runs[0]
r2.font.size      = Pt(18)
r2.font.color.rgb = GOLD_LIGHT
r2.font.italic    = True

p_div = cc.add_paragraph('─' * 50)
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_div.paragraph_format.space_before = Pt(6)
p_div.paragraph_format.space_after  = Pt(6)
p_div.runs[0].font.color.rgb = GOLD

p_title = cc.add_paragraph('GUÍA DE USO — CRM & DASHBOARD')
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(4)
r3 = p_title.runs[0]
r3.font.size      = Pt(16)
r3.font.color.rgb = WHITE
r3.bold           = True

p_info = cc.add_paragraph('Manual completo del sistema para el equipo de Kairos')
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_info.paragraph_format.space_before = Pt(0)
p_info.paragraph_format.space_after  = Pt(40)
p_info.runs[0].font.size      = Pt(10)
p_info.runs[0].font.color.rgb = GOLD_LIGHT
p_info.runs[0].italic         = True

# Index
idx = doc.add_paragraph()
idx.alignment = WD_ALIGN_PARAGRAPH.CENTER
idx.paragraph_format.space_before = Pt(16)
r_idx = idx.add_run('ÍNDICE DE SECCIONES')
r_idx.bold           = True
r_idx.font.size      = Pt(12)
r_idx.font.color.rgb = BROWN_DARK

sections_list = [
    ('0', 'Introducción y Primeros Pasos',  'Cómo acceder, navegar y entender el sistema desde cero'),
    ('1', 'Dashboard',               'Resumen ejecutivo y métricas clave'),
    ('2', 'Leads',                   'Base de prospectos, detalle del contacto, tareas y notas'),
    ('3', 'Mayoristas',              'Gestión de distribuidores y proveedores'),
    ('4', 'Pipeline de Ventas',      'Tablero Kanban del embudo comercial'),
    ('5', 'Campañas',                'Crear campañas, enviar emails, métricas y seguimiento WA'),
    ('6', 'Órdenes',                 'Gestión de pedidos de principio a fin'),
    ('7', 'Catálogo de Productos',   'Productos, precios, stock, Google Sheets auto-sync'),
    ('8', 'Scraper de Leads',        'Búsqueda automática y enriquecimiento de contactos'),
    ('✦', 'Glosario',               'Definición de todos los términos del sistema'),
]

for num, name, desc in sections_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_n  = p.add_run(f'  {num}.  ')
    r_n.bold            = True
    r_n.font.color.rgb  = GOLD
    r_n.font.size       = Pt(10.5)
    r_nm = p.add_run(name)
    r_nm.bold           = True
    r_nm.font.size      = Pt(10.5)
    r_nm.font.color.rgb = BROWN_DARK
    r_de = p.add_run(f'  ·  {desc}')
    r_de.font.size      = Pt(9.5)
    r_de.font.color.rgb = BROWN_MID

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════
#  SECTION 0 — INTRODUCCIÓN Y PRIMEROS PASOS
# ══════════════════════════════════════════════════════════════════════════

section_header(doc, '0', 'INTRODUCCIÓN Y PRIMEROS PASOS',
               'Cómo acceder al sistema, qué es cada sección y por dónde empezar')

h2(doc, '¿Qué es el sistema Kairos CRM?')

body(doc, 'Kairos CRM es el sistema de gestión comercial de Kairos Distribuidora. '
         'Concentra en un solo lugar todo lo necesario para conseguir clientes nuevos, '
         'comunicarse con ellos, hacer pedidos y mantener el catálogo de productos actualizado. '
         'No hace falta instalar nada — funciona desde cualquier navegador web.')

h2(doc, 'Cómo acceder')

data_table(doc,
    ['Qué', 'Dónde'],
    [
        ('URL del sistema',    'kairos.polkorp.com'),
        ('Catálogo público',   'kairos.polkorp.com/public/catalog (accesible por clientes sin login)'),
        ('Compatibilidad',     'Chrome, Edge, Firefox, Safari — en computadora, tablet o celular'),
    ],
    col_widths=[4.0, 11.0]
)

tip_box(doc, '💡  Guardá kairos.polkorp.com como favorito en tu navegador para acceder rápido cada día.',
        color=GOLD)

h2(doc, 'El menú lateral — cómo navegar el sistema')

body(doc, 'Al ingresar al sistema verás una barra de navegación lateral en el lado izquierdo. '
         'Desde ahí podés ir a cualquier sección con un click:')

data_table(doc,
    ['Ícono en el menú', 'Sección', 'Para qué sirve'],
    [
        ('🏠 Casa',           'Dashboard',    'Pantalla de inicio — resumen del negocio en tiempo real'),
        ('👥 Personas',        'Leads',        'Lista completa de prospectos y clientes minoristas'),
        ('🏢 Edificio',        'Mayoristas',   'Lista de distribuidores y mayoristas'),
        ('📊 Barras',          'Pipeline',     'Vista visual del embudo de ventas en tablero Kanban'),
        ('📣 Megáfono',        'Campañas',     'Crear y enviar emails o mensajes masivos'),
        ('📦 Caja',            'Órdenes',      'Pedidos de clientes — desde confirmación hasta entrega'),
        ('🛍 Bolsa',           'Catálogo',     'Productos, precios, stock y sincronización'),
        ('🔍 Lupa',            'Scraper',      'Búsqueda automática de prospectos nuevos'),
    ],
    col_widths=[3.0, 3.0, 9.0]
)

h2(doc, '¿Por dónde empezar si es la primera vez?')

body(doc, 'Si estás usando el sistema por primera vez, seguí este orden:')

numbered(doc, 'Sección 8 — Scraper: lanzá el scraper para llenar la base con prospectos nuevos.')
numbered(doc, 'Sección 8 — Scraper: corrí el Enriquecedor para conseguir emails de los prospectos.')
numbered(doc, 'Sección 7 — Catálogo: cargá tus productos (manualmente o importando desde Kairosdis).')
numbered(doc, 'Sección 7 — Catálogo: conectá tu planilla de Google Sheets para mantener los precios.')
numbered(doc, 'Sección 5 — Campañas: creá tu primera campaña y enviala al segmento que elijas.')
numbered(doc, 'Sección 1 — Dashboard: revisá el resumen a diario para estar al tanto del negocio.')

tip_box(doc, '⚠️  Para enviar emails desde el sistema es necesario tener Brevo (ex Sendinblue) configurado. '
             'Si los emails no llegan, contactá al administrador del sistema para verificar la configuración.',
        color=AMBER)

h2(doc, 'Conceptos básicos que vas a ver en todo el sistema')

body(doc, 'Estos términos aparecen en múltiples secciones — es importante entenderlos desde el comienzo:')

data_table(doc,
    ['Término', 'Qué significa en el sistema'],
    [
        ('Lead',         'Un negocio o persona que podría comprar productos Kairos. Puede ser una tienda, un emprendedor, etc.'),
        ('Estado',       'La etapa en la que está un lead en el proceso de venta: Nuevo → Contactado → Interesado → Cliente → Descartado.'),
        ('Score IA',     'Calificación automática del 0 al 10 que el sistema le da a cada lead según cuántos datos tiene cargados.'),
        ('Campaña',      'Un mensaje masivo enviado a un grupo de leads. Puede ser un email o una lista de links de WhatsApp.'),
        ('Orden',        'Un pedido de compra confirmado de un cliente. Incluye productos, cantidades y precio total.'),
        ('Catálogo',     'La lista completa de productos disponibles con precio y stock. Tiene una versión pública para clientes.'),
        ('Enriquecer',   'Proceso automático que visita sitios web de los leads para encontrar emails y datos de contacto faltantes.'),
        ('Pipeline',     'Vista Kanban del embudo de ventas — los leads se mueven entre columnas según avanzan en el proceso.'),
    ],
    col_widths=[3.0, 12.0]
)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════
#  SECTION 1 — DASHBOARD
# ══════════════════════════════════════════════════════════════════════════

section_header(doc, '1', 'DASHBOARD',
               'Vista de inicio — Resumen general del negocio')

body(doc, 'El Dashboard es la primera pantalla al entrar al sistema. Muestra de un vistazo el estado '
         'completo del negocio: leads, revenue, tareas pendientes y productos.')

h2(doc, 'Tarjetas de métricas rápidas')

data_table(doc,
    ['Tarjeta', 'Qué muestra', 'Link directo'],
    [
        ('Total Leads',       'Cantidad total de contactos en la base',            'Sección Leads'),
        ('Con Email',         'Leads con email registrado — listos para campañas', 'Sección Leads'),
        ('Sin Email',         'Leads que necesitan enriquecimiento (rojo si >500)', 'Scraper'),
        ('Clientes',          'Leads que ya realizaron al menos un pedido',         'Sección Leads'),
        ('Revenue del Mes',   'Facturación acumulada en ARS del mes en curso',      'Órdenes'),
        ('Tareas Vencidas',   'Seguimientos pendientes — rojo si hay alguno',       'Leads'),
        ('Productos',         'Total de productos en el catálogo',                  'Catálogo'),
    ],
    col_widths=[3.5, 7.5, 3.5]
)

h2(doc, 'Gráficos disponibles')

bullet(doc, 'Leads por Provincia (barras) — top 8 provincias con más prospectos.')
bullet(doc, 'Leads por Estado (torta) — distribución del embudo en porcentajes.')
bullet(doc, 'Embudo de Leads — barra de progreso por etapa con conteo y %.')
bullet(doc, 'Últimos Leads — lista de los 8 más recientes con acceso al detalle.')
bullet(doc, 'Órdenes por Mes — línea de actividad comercial mensual.')
bullet(doc, 'Revenue por Mes — barras de facturación en ARS.')

h2(doc, 'Importar catálogo desde Kairosdis')

body(doc, 'En la parte inferior del Dashboard hay una preview del catálogo con el botón '
         '"Importar Kairosdis". Al presionarlo, el sistema descarga automáticamente todos los '
         'productos de kairosdis.com.ar y actualiza el catálogo en tiempo real.')

tip_box(doc, '💡  Hacé click en cualquier tarjeta del Dashboard para ir directamente a la '
             'sección correspondiente.', color=GOLD)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════
#  SECTION 2 — LEADS
# ══════════════════════════════════════════════════════════════════════════

section_header(doc, '2', 'LEADS',
               'Base de datos principal de prospectos y clientes minoristas')

body(doc, 'Cada lead representa una tienda, local o emprendedor que puede comprar productos Kairos. '
         'Esta sección es el corazón del CRM: es donde está toda la información de tus contactos '
         'y desde donde se inician las comunicaciones.')

tip_box(doc, '💡  ¿Cómo se agregan leads? Los leads entran al sistema de dos formas: '
             '(1) automáticamente desde el Scraper, que busca negocios en Google Places; '
             '(2) importando un archivo CSV con los datos. No hay un formulario para agregar '
             'un lead de a uno — si necesitás agregar un contacto individual, usá la importación '
             'con un CSV de una sola fila.',
        color=BLUE)

h2(doc, 'Score IA — sistema de calificación automática')

data_table(doc,
    ['Score', 'Color', 'Significado'],
    [
        (('7 – 10', True, GREEN),  ('Verde',    False, GREEN),  'Lead completo: tiene teléfono, email, web y buen rating'),
        (('4 – 6',  True, AMBER),  ('Amarillo', False, AMBER),  'Lead parcial: falta enriquecer algún dato'),
        (('0 – 3',  True, RED),    ('Rojo',     False, RED),    'Lead muy incompleto: sólo tiene nombre'),
    ],
    col_widths=[2.5, 3.0, 9.0]
)

h2(doc, 'Filtros disponibles')

data_table(doc,
    ['Filtro', 'Cómo usarlo'],
    [
        ('Buscar empresa',    'Escribí el nombre — filtra en tiempo real'),
        ('Provincia',         'Desplegable con todas las provincias presentes en la base'),
        ('Rubro',             'Tipo de negocio: tienda holística, sahumerios, santería, etc.'),
        ('Estado',            'Nuevo / Contactado / Interesado / Cliente / Descartado'),
        ('Solo con email',    'Toggle — ideal antes de lanzar una campaña de email'),
        ('Solo con teléfono', 'Toggle — filtra los que tienen número para WhatsApp'),
    ],
    col_widths=[4.5, 10.5]
)

h2(doc, 'Estados del lead')

status_badge_table(doc, [
    ('  Nuevo  ', 'gray'),
    ('  Contactado  ', 'blue'),
    ('  Interesado  ', 'amber'),
    ('  Cliente  ', 'green'),
    ('  Descartado  ', 'red'),
])

body(doc, 'El estado se puede cambiar directamente desde la lista sin abrir el detalle: '
         'hacé click en el desplegable de la columna "Estado" en la fila del lead.')

h2(doc, 'Acciones masivas — barra flotante')

body(doc, 'Seleccioná uno o más leads con el checkbox. Aparece automáticamente una barra '
         'flotante en la parte inferior de la pantalla con estas acciones:')

data_table(doc,
    ['Acción', 'Descripción', 'Requiere'],
    [
        ('Email',     'Envía un email a los seleccionados. Podés generar el texto con IA.', 'Email configurado (Brevo)'),
        ('WhatsApp',  'Genera links wa.me para abrir cada chat directo.',                   'Lead con teléfono'),
        ('Catálogo',  'Envía el link al catálogo público por email.',                       'Lead con email'),
    ],
    col_widths=[3.0, 8.5, 4.5]
)

h2(doc, 'Paso a paso: enviar email con IA')

numbered(doc, 'Seleccioná los leads deseados con el checkbox.')
numbered(doc, 'Hacé click en "Email" en la barra flotante.')
numbered(doc, 'Presioná "Generar con IA en español" — el sistema redacta asunto y cuerpo automáticamente.')
numbered(doc, 'Editá el texto si querés. Usá {empresa} para personalizar con el nombre de cada lead.')
numbered(doc, 'Hacé click en "Enviar".')

h2(doc, 'Paso a paso: generar links de WhatsApp')

numbered(doc, 'Seleccioná leads que tengan teléfono registrado.')
numbered(doc, 'Hacé click en "WhatsApp" en la barra flotante.')
numbered(doc, 'Escribí el mensaje.')
numbered(doc, 'Hacé click en "Generar Links".')
numbered(doc, 'Hacé click en cada link para abrir el chat directamente en WhatsApp.')

h2(doc, 'Importar y exportar CSV')

data_table(doc,
    ['Operación', 'Cómo', 'Notas'],
    [
        ('Exportar CSV', 'Botón "Exportar CSV" — arriba a la derecha', 'Exporta exactamente los leads filtrados en ese momento'),
        ('Importar CSV', 'Botón "Importar CSV" — sube un archivo .csv', 'Columnas: empresa, telefono, email, ciudad, provincia, rubro, website'),
    ],
    col_widths=[3.5, 5.5, 6.5]
)

tip_box(doc, '⚠️  Al importar, los duplicados (mismo nombre de empresa) se omiten automáticamente. '
             'El archivo debe tener encabezados en la primera fila y acepta CSV de Excel (UTF-8 o BOM).',
        color=AMBER)

h2(doc, 'Crear una lista con precios para cotizar')

body(doc, 'Para armar una lista de contactos con referencia de precios y enviarla a un segmento específico:')

numbered(doc, 'En Leads, aplicá los filtros que necesitás: Provincia, Rubro, Estado = "Interesado", activá "Solo con email".')
numbered(doc, 'Hacé click en "Exportar CSV" — el archivo descargado tiene nombre, empresa, teléfono, email, provincia, rubro.')
numbered(doc, 'Abrí el CSV en Excel o Google Sheets.')
numbered(doc, 'En el Catálogo, exportá también el CSV de productos (tiene Nombre, Precio Minorista, Precio Mayorista, Stock).')
numbered(doc, 'Combiná ambas planillas según necesitás: copia las columnas de precios relevantes junto a los contactos.')
numbered(doc, 'Usá la tabla resultante para enviar cotizaciones o como referencia para llamadas comerciales.')

tip_box(doc, '💡  Alternativa rápida: desde la barra flotante de Leads, seleccioná los contactos, '
             'hacé click en "Catálogo" para enviarles el link al catálogo público con todos los precios actualizados. '
             'No requiere exportar nada.', color=GOLD)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════
#  SECTION 2B — DETALLE DEL LEAD
# ══════════════════════════════════════════════════════════════════════════

section_header(doc, '2b', 'DETALLE DEL LEAD',
               'Ficha completa del contacto — notas, tareas y seguimiento')

body(doc, 'Hacé click en el nombre de cualquier lead para abrir su ficha completa. '
         'Esta pantalla centraliza toda la información y el historial del contacto.')

h2(doc, 'Panel de información de contacto')

body(doc, 'En la parte superior izquierda se muestra y puede editar toda la información del negocio:')

data_table(doc,
    ['Campo', 'Descripción', 'Editable'],
    [
        ('Empresa',    'Nombre del local o emprendimiento',                        'Sí'),
        ('Teléfono',   'Número con código de área (sin 0, sin 15)',                'Sí'),
        ('Email',      'Dirección de email del contacto',                         'Sí'),
        ('Ciudad',     'Ciudad donde está el negocio',                            'Sí'),
        ('Provincia',  'Provincia — usada por los filtros y campañas',            'Sí'),
        ('Rubro',      'Tipo de negocio: holístico, sahumerios, santería, etc.',  'Sí'),
        ('Website',    'URL del sitio web — el enriquecedor extrae datos de aquí','Sí'),
        ('Rating',     'Calificación de Google Places (1-5 estrellas)',           'Solo lectura'),
        ('Dirección',  'Dirección física del negocio',                            'Sí'),
    ],
    col_widths=[3.0, 8.5, 2.5]
)

h2(doc, 'Panel de estado y calificación')

body(doc, 'El panel derecho superior permite gestionar la situación comercial del lead:')

bullet(doc, 'Estado: desplegable con las 5 etapas (Nuevo / Contactado / Interesado / Cliente / Descartado).')
bullet(doc, 'Score IA: puntaje automático calculado según los datos disponibles (no editable).')
bullet(doc, 'Observaciones: campo de texto libre para notas internas sobre el contacto.')

tip_box(doc, '💡  Las Observaciones son privadas del equipo — nunca se envían al contacto. '
             'Usalas para registrar acuerdos verbales, preferencias de productos, mejor horario para llamar, etc.',
        color=GOLD)

h2(doc, 'Timeline de notas — actividad del contacto')

body(doc, 'La sección central muestra todas las interacciones en orden cronológico:')

bullet(doc, 'Cada entrada tiene fecha, hora y texto de la nota.')
bullet(doc, 'Para agregar una nota: escribí en el campo "Nueva nota..." y hacé click en "Agregar Nota".')
bullet(doc, 'Las notas son permanentes — útil para registrar llamadas, reuniones, respuestas de email.')

h2(doc, 'Tareas y seguimiento — paso a paso')

body(doc, 'La sección de tareas permite organizar el seguimiento futuro de cada lead:')

numbered(doc, 'Hacé click en "Nueva Tarea".')
numbered(doc, 'Escribí el título de la tarea (ej: "Llamar para confirmar pedido", "Enviar catálogo actualizado").')
numbered(doc, 'Opcionalmente agregá una descripción con más detalle.')
numbered(doc, 'Establecé una fecha de vencimiento usando el calendario.')
numbered(doc, 'Hacé click en "Crear" — la tarea aparece en la lista.')
numbered(doc, 'Cuando la completés, hacé click en el checkbox para marcarla como hecha.')

data_table(doc,
    ['Estado de tarea', 'Cómo se muestra'],
    [
        (('Pendiente',  True, AMBER), 'Fondo normal, sin tachado'),
        (('Vencida',    True, RED),   'Fondo rojo — la fecha límite ya pasó'),
        (('Completada', True, GREEN), 'Tachada y con checkmark verde'),
    ],
    col_widths=[4.0, 11.0]
)

tip_box(doc, '⚠️  Las tareas vencidas aparecen en el Dashboard como "Tareas Vencidas" con contador rojo. '
             'Revisalas a diario para no perder seguimientos importantes.',
        color=AMBER)

h2(doc, 'Historial de órdenes del lead')

body(doc, 'Al final de la ficha se listan todas las órdenes que hizo ese cliente: '
         'número de orden, fecha, estado y monto total en ARS. '
         'Hacé click en una orden para ver el detalle completo de productos y cantidades.')

tip_box(doc, '💡  Desde esta pantalla también podés crear una nueva orden directamente para este cliente '
             'sin volver al listado de Órdenes — ahorra tiempo cuando el cliente llama para hacer un pedido.',
        color=GOLD)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════
#  SECTION 3 — MAYORISTAS
# ══════════════════════════════════════════════════════════════════════════

section_header(doc, '3', 'MAYORISTAS',
               'Distribuidores, importadores y proveedores')

body(doc, 'Funciona igual que la sección Leads pero filtra únicamente contactos de tipo mayorista. '
         'La lógica de estados, score IA y acciones masivas es exactamente la misma.')

h2(doc, 'Diferencias clave vs. Leads')

data_table(doc,
    ['Aspecto', 'Leads (minoristas)', 'Mayoristas'],
    [
        ('Tipo de contacto', 'Tiendas y locales de venta al público', 'Distribuidores, importadores, mayoristas'),
        ('Queries del Scraper', 'Sahumerios, tiendas holísticas, etc.', '"Distribuidor sahumerios Argentina", "Mayorista holísticos", etc.'),
        ('Precio de referencia', 'Precio minorista', 'Precio mayorista'),
    ],
    col_widths=[4.0, 5.5, 5.5]
)

tip_box(doc, '💡  Para agregar mayoristas a la base, lanzá el scraper desde esta sección '
             'con el botón "Iniciar Scraper". Usa queries especializadas en distribuidores.',
        color=GOLD)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════
#  SECTION 4 — PIPELINE
# ══════════════════════════════════════════════════════════════════════════

section_header(doc, '4', 'PIPELINE DE VENTAS',
               'Tablero Kanban visual del embudo comercial')

body(doc, 'Vista Kanban con todos los leads organizados en columnas por estado. '
         'Es la forma más visual de gestionar el avance de cada prospecto.')

h2(doc, 'Las 5 columnas del Pipeline')

data_table(doc,
    ['Columna', 'Color', 'Descripción'],
    [
        (('Nuevo',       True, RGBColor(0x64,0x74,0x8B)), ('Gris',     False, RGBColor(0x64,0x74,0x8B)), 'Lead recién ingresado, sin contacto previo'),
        (('Contactado',  True, BLUE),                     ('Azul',     False, BLUE),                     'Ya recibió un primer contacto (email o llamada)'),
        (('Interesado',  True, AMBER),                    ('Amarillo', False, AMBER),                    'Está en conversación activa'),
        (('Cliente',     True, GREEN),                    ('Verde',    False, GREEN),                    'Realizó al menos una compra confirmada'),
        (('Descartado',  True, RED),                      ('Rojo',     False, RED),                      'Fuera del embudo de ventas'),
    ],
    col_widths=[3.0, 2.5, 9.5]
)

h2(doc, 'Cómo mover leads entre columnas')

body(doc, 'Cada tarjeta de lead tiene tres botones de acción en la parte inferior:')

data_table(doc,
    ['Botón', 'Acción'],
    [
        ('◀  (flecha izquierda)', 'Retrocede el lead a la etapa anterior'),
        ('▶  (flecha derecha)',   'Avanza el lead a la siguiente etapa'),
        ('✕  (X roja)',           'Descarta el lead directamente (va a columna Descartado)'),
    ],
    col_widths=[5.0, 10.0]
)

tip_box(doc, '💡  También podés hacer click en la tarjeta para abrir el detalle completo del lead '
             'y editar cualquier dato o agregar notas de seguimiento.',
        color=GOLD)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════
#  SECTION 5 — CAMPAÑAS
# ══════════════════════════════════════════════════════════════════════════

section_header(doc, '5', 'CAMPAÑAS',
               'Email marketing, creación paso a paso, métricas y seguimiento por WhatsApp')

body(doc, 'Una campaña es un mensaje enviado al mismo tiempo a un grupo de leads. '
         'Podés elegir exactamente a quién le llega (por provincia, rubro o estado), '
         'escribir el texto vos o usar la IA para redactarlo, y ver después '
         'cuántos lo abrieron, cuántos hicieron click y cuántos se convirtieron en clientes.')

tip_box(doc, '⚠️  Requisito previo: para enviar campañas de email, el sistema necesita '
             'tener Brevo (ex Sendinblue) configurado con una API key válida. '
             'Si los emails no llegan, consultá al administrador. Las campañas de WhatsApp '
             'no requieren ninguna configuración especial — solo generan links que vos abrís.',
        color=AMBER)

h2(doc, 'Panel de métricas de campañas')

data_table(doc,
    ['Indicador', 'Qué mide'],
    [
        ('Total Campañas',          'Cantidad total de campañas creadas'),
        ('Emails Enviados',         'Volumen acumulado de todos los envíos'),
        ('Tasa de Apertura Prom.',  'Promedio de apertura (%) entre todas las campañas'),
        ('Conversiones',            'Leads que pasaron a estado "cliente" tras una campaña'),
    ],
    col_widths=[5.5, 9.5]
)

h2(doc, 'Estados de una campaña')

status_badge_table(doc, [
    ('  Borrador  ', 'gray'),
    ('  Programado  ', 'blue'),
    ('  Activo  ', 'green'),
    ('  Pausado  ', 'amber'),
    ('  Enviado  ', 'green'),
    ('  Cancelado  ', 'red'),
])

h2(doc, 'Crear una nueva campaña — wizard de 3 pasos')

body(doc, 'Hacé click en "Nueva Campaña" (botón superior derecho). Se abre un asistente de 3 pasos:')

doc.add_paragraph()

wizard_step_box(doc, 1, 'Configurar campaña',
    'Nombre, tipo de envío y segmento de destinatarios')

body(doc, 'En el primer paso configurás a quién le llega la campaña:')

data_table(doc,
    ['Campo', 'Opciones / Descripción'],
    [
        ('Nombre de campaña',   'Texto libre — solo visible en el CRM, no lo ve el destinatario'),
        ('Tipo',                'Email  o  WhatsApp (genera links wa.me en lugar de emails)'),
        ('Provincia',           'Filtrar solo leads de una provincia específica (opcional)'),
        ('Rubro',               'Filtrar por tipo de negocio: tienda holística, sahumerios, etc. (opcional)'),
        ('Estado del lead',     'Nuevo / Contactado / Interesado / Cliente (opcional)'),
        ('Solo con email',      'Toggle — si está activo, excluye leads sin email registrado'),
    ],
    col_widths=[4.0, 11.0]
)

tip_box(doc, '💡  Si dejás todos los filtros vacíos, la campaña llega a todos los leads de la base. '
             'Combiná filtros para segmentar con precisión: por ejemplo, "solo leads de Buenos Aires '
             'que estén en estado Interesado y tengan email".', color=GOLD)

doc.add_paragraph()

wizard_step_box(doc, 2, 'Contenido del mensaje',
    'Redactá el asunto y el cuerpo, o generá el texto con inteligencia artificial')

body(doc, 'En el segundo paso escribís o generás el texto de la campaña:')

bullet(doc, 'Asunto: línea de asunto del email (no aplica para WhatsApp).')
bullet(doc, 'Cuerpo: texto del mensaje. Podés usar variables de personalización.')
bullet(doc, 'Generá con IA: el sistema redacta asunto y cuerpo en español automáticamente según '
            'el tipo de campaña y el segmento elegido.')

data_table(doc,
    ['Variable', 'Se reemplaza por', 'Ejemplo en el texto'],
    [
        ('{nombre}',  'Nombre del contacto (si está registrado)', '"Hola {nombre}, te contactamos desde Kairos..."'),
        ('{empresa}', 'Nombre del negocio o empresa del lead',    '"Escribimos a {empresa} para presentarles..."'),
    ],
    col_widths=[3.0, 5.5, 7.0]
)

tip_box(doc, '💡  Podés mezclar texto generado con IA y edición manual. Generá primero y después '
             'ajustá el tono o agregá información específica del producto que querés destacar.',
        color=GOLD)

doc.add_paragraph()

wizard_step_box(doc, 3, 'Revisión y envío',
    'Revisá el resumen, confirmá cuántos leads recibirán la campaña y enviá')

body(doc, 'El tercer paso muestra un resumen antes de confirmar:')

bullet(doc, 'Nombre de la campaña y tipo de envío.')
bullet(doc, 'Contador de leads que cumplieron los filtros — sabés exactamente a cuántas personas llega.')
bullet(doc, 'Preview del asunto y los primeros caracteres del cuerpo.')
bullet(doc, '"Enviar Campaña": confirma y lanza el envío inmediatamente.')

tip_box(doc, '⚠️  Una vez enviada, la campaña no se puede cancelar ni modificar. '
             'Verificá el contador de destinatarios y el preview del texto antes de confirmar.',
        color=AMBER)

h2(doc, 'Ver el detalle y métricas de una campaña enviada')

body(doc, 'Hacé click en el nombre de cualquier campaña para abrir su página de métricas:')

data_table(doc,
    ['Indicador', 'Qué mide', 'Cómo interpretar'],
    [
        (('Enviados',     True, BROWN_DARK), 'Total de emails despachados al servidor', 'Siempre igual al contador de leads al momento del envío'),
        (('Abiertos',     True, BLUE),       'Emails que el destinatario abrió',         'Tasa de apertura normal: 20-40%'),
        (('Clicks',       True, AMBER),      'Clicks en links dentro del email',         'Indica interés real en el contenido'),
        (('Respondidos',  True, GREEN),      'Respuestas directas al email',             'Alto valor — requieren seguimiento personal'),
        (('Convertidos',  True, GREEN),      'Leads que pasaron a estado Cliente',       'Métrica final de efectividad de la campaña'),
    ],
    col_widths=[3.0, 5.5, 6.5]
)

body(doc, 'Debajo de las métricas se muestra la tabla individual de envíos: '
         'cada lead con su estado (enviado / abierto / click / respondido) y la fecha de cada acción.')

tip_box(doc, '💡  Combiná la vista de métricas con "Seguimiento WA": usá los leads que abrieron '
             'el email pero no respondieron para un seguimiento personalizado por WhatsApp.',
        color=GREEN)

h2(doc, 'Duplicar una campaña')

body(doc, 'El ícono de copiar (🗐) en la columna Acciones duplica la campaña con todos sus '
         'parámetros. Útil para reutilizar una campaña exitosa cambiando sólo el segmento '
         'o la fecha de envío.')

h2(doc, 'Seguimiento por WhatsApp — "Seguimiento WA"')

body(doc, 'El botón verde "Seguimiento WA" muestra automáticamente todos los leads a los que '
         'se les envió un email hace más de 3 días y todavía no respondieron.')

numbered(doc, 'Hacé click en "Seguimiento WA".')
numbered(doc, 'Se carga la lista de leads pendientes con sus teléfonos.')
numbered(doc, 'Hacé click en cada link para abrir el chat de WhatsApp directamente.')

tip_box(doc, '💡  El seguimiento WA es ideal para convertir leads "contactados" en "interesados": '
             'un mensaje personal de seguimiento por WhatsApp suele tener mucho más impacto que un '
             'segundo email.',
        color=GREEN)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════
#  SECTION 6 — ÓRDENES
# ══════════════════════════════════════════════════════════════════════════

section_header(doc, '6', 'ÓRDENES',
               'Gestión de pedidos desde la confirmación hasta la entrega')

body(doc, 'Vista Kanban con 5 etapas que representa el ciclo completo de un pedido.')

h2(doc, 'Etapas de una orden')

data_table(doc,
    ['Etapa', 'Color', 'Descripción'],
    [
        (('Borrador',       True, RGBColor(0x64,0x74,0x8B)), ('Gris',    False, RGBColor(0x64,0x74,0x8B)), 'Orden cargada pero sin confirmar aún'),
        (('Confirmado',     True, BLUE),                     ('Azul',    False, BLUE),                     'El cliente confirmó el pedido'),
        (('En Preparación', True, AMBER),                    ('Amarillo',False, AMBER),                    'El pedido está siendo armado/empacado'),
        (('Despachado',     True, RGBColor(0x76,0x29,0xD8)), ('Violeta', False, RGBColor(0x76,0x29,0xD8)),'Fue enviado al cliente'),
        (('Entregado',      True, GREEN),                    ('Verde',   False, GREEN),                    'Recibido y cerrado'),
    ],
    col_widths=[3.5, 2.5, 9.0]
)

h2(doc, 'Crear una nueva orden — paso a paso')

numbered(doc, 'Hacé click en "Nueva Orden" (botón superior derecho).')
numbered(doc, 'Seleccioná el cliente desde el desplegable (busca dentro de todos los leads).')
numbered(doc, 'Seleccioná el primer producto y la cantidad.')
numbered(doc, 'Hacé click en "+ Agregar producto" para sumar más ítems a la misma orden.')
numbered(doc, 'El precio mayorista se muestra al lado de cada producto como referencia.')
numbered(doc, 'Hacé click en "Crear Orden" — aparece automáticamente en columna Borrador.')

h2(doc, 'Ver y avanzar el estado de una orden')

body(doc, 'Hacé click en la tarjeta de la orden para abrir el detalle. Desde ahí podés:')

bullet(doc, 'Ver el resumen de productos, cantidades y monto total en ARS.')
bullet(doc, 'Cambiar la etapa (Borrador → Confirmado → En Preparación → Despachado → Entregado).')
bullet(doc, 'Agregar notas o comentarios sobre el pedido.')

tip_box(doc, '💡  Desde el detalle de un lead podés ir directamente a "Ver sus órdenes" '
             'o crear una nueva orden para ese cliente sin volver al listado.',
        color=GOLD)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════
#  SECTION 7 — CATÁLOGO
# ══════════════════════════════════════════════════════════════════════════

section_header(doc, '7', 'CATÁLOGO DE PRODUCTOS',
               'Productos, precios, stock, imágenes y sincronización automática con Google Sheets')

h2(doc, 'Agregar un producto manualmente')

numbered(doc, 'Hacé click en "Agregar Producto" (botón superior derecho).')
numbered(doc, 'Subí la imagen arrastrando un archivo o haciendo click en el área. '
              'Se redimensiona automáticamente a máx. 800px.')
numbered(doc, 'Completá Nombre (obligatorio), Categoría, Descripción.')
numbered(doc, 'Ingresá Precio Minorista (precio al público), Precio Mayorista y Stock.')
numbered(doc, 'Activá el toggle "Activo" para que aparezca en el catálogo público.')
numbered(doc, 'Activá "Destacado" (⭐) para que aparezca primero en la grilla.')
numbered(doc, 'Hacé click en "Crear".')

h2(doc, 'Editar un producto o cambiar precios')

numbered(doc, 'Hacé click en el ícono de lápiz (✏️) en la tarjeta del producto.')
numbered(doc, 'Modificá los campos necesarios: nombre, precios, stock, imagen, etc.')
numbered(doc, 'Hacé click en "Guardar".')

h2(doc, 'Categorías disponibles')

cats = [
    'Aceites esenciales', 'Aceites portadores', 'Hidrolatos', 'Blends',
    'Difusores', 'Nebulizadores', 'Accesorios', 'Cristales', 'Libros', 'Kits',
]
# Show as a 2-col table
t = doc.add_table(rows=(len(cats) + 1) // 2, cols=2)
no_borders(t)
for i, cat in enumerate(cats):
    r_idx = i // 2
    c_idx = i % 2
    cell = t.cell(r_idx, c_idx)
    set_cell_bg(cell, GOLD_LIGHT if (r_idx + c_idx) % 2 == 0 else WHITE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.3)
    bullet_r = p.add_run('▸  ' + cat)
    bullet_r.font.size      = Pt(9.5)
    bullet_r.font.color.rgb = BROWN_DARK
doc.add_paragraph()

h2(doc, 'Importar Kairosdis')

body(doc, 'El botón "Importar Kairosdis" scrapea kairosdis.com.ar y sincroniza todos los productos '
         '(nombre, precio, imágenes). Muestra porcentaje de avance en tiempo real. '
         'Al terminar informa cuántos productos son nuevos y cuántos fueron actualizados.')

h2(doc, 'Google Sheets como catálogo maestro — sincronización automática')

body(doc, 'La función más potente del catálogo: conectar una planilla de Google Sheets para que '
         'sea la fuente de verdad. Cuando actualizás un precio o stock en la planilla, '
         'el sistema lo refleja en el CRM automáticamente sin ningún paso adicional.')

data_table(doc,
    ['Característica', 'Detalle'],
    [
        ('Sync automático',     'El sistema verifica la planilla cada 10 minutos y aplica los cambios'),
        ('Sin intervención',    'No hace falta tocar el CRM — solo editá la planilla y esperá máx. 10 min'),
        ('Creación de nuevos',  'Si agregás una fila nueva en la planilla, se crea el producto en el CRM'),
        ('Matching inteligente','El sistema busca por ID primero; si no lo encuentra, busca por nombre del producto'),
        ('Último sync visible', 'La pantalla muestra la hora del último sync y cuántos productos fueron creados/actualizados'),
    ],
    col_widths=[4.0, 11.0]
)

h2(doc, 'Conectar la planilla de Google Sheets — primera vez')

numbered(doc, 'En el Catálogo, hacé click en "Exportar CSV" — descarga todos los productos actuales.')
numbered(doc, 'Abrí Google Sheets en drive.google.com. Creá una nueva hoja en blanco.')
numbered(doc, 'Importá el CSV: Archivo → Importar → Subir → elegí el archivo descargado.')
numbered(doc, 'Seleccioná "Reemplazar hoja de cálculo" y aceptá. Ahora tenés todos los productos con encabezados.')
numbered(doc, 'Compartí la planilla como pública: Compartir → Cambiar a "Cualquiera con el link" → rol Lector.')
numbered(doc, 'Copiá el link de la planilla (debe tener el formato: docs.google.com/spreadsheets/d/ID/...).')
numbered(doc, 'En el CRM, Catálogo, hacé click en "Conectar Google Sheet" (botón dorado).')
numbered(doc, 'Pegá el link (o el ID) de la planilla y hacé click en "Guardar y sincronizar".')
numbered(doc, 'El sistema hace una sincronización inmediata y muestra los resultados. '
              'Desde ahora synca automáticamente cada 10 minutos.')

tip_box(doc, '✅  Una vez conectada, el botón en el Catálogo se pone verde con un checkmark y dice '
             '"Sheet conectado". Abajo aparece el último horario de sync y cuántos productos '
             'fueron creados o actualizados en esa pasada.', color=GREEN)

h2(doc, 'Columnas que podés editar en la planilla')

data_table(doc,
    ['Columna', 'Qué controla', 'Editable'],
    [
        (('ID',               True,  BROWN_DARK), 'Identificador único del producto',           ('NO — nunca modificar', True, RED)),
        (('Precio Minorista', True,  BROWN_DARK), 'Precio de venta al público (ARS)',           ('Sí', True, GREEN)),
        (('Precio Mayorista', True,  BROWN_DARK), 'Precio de venta a revendedores (ARS)',       ('Sí', True, GREEN)),
        (('Precio Promo',     True,  BROWN_DARK), 'Precio promocional opcional (ARS)',          ('Sí', True, GREEN)),
        (('Stock',            True,  BROWN_DARK), 'Unidades disponibles (número entero)',       ('Sí', True, GREEN)),
        (('Activo',           True,  BROWN_DARK), 'Sí = visible en catálogo  /  No = oculto',  ('Sí', True, GREEN)),
        (('Nombre',           False, BROWN_MID),  'Al agregar fila nueva, usa este nombre',     ('Sí para nuevos', False, AMBER)),
        (('Descripcion',      False, BROWN_MID),  'Al agregar fila nueva, usa esta descripción',('Sí para nuevos', False, AMBER)),
    ],
    col_widths=[3.5, 7.5, 4.0]
)

tip_box(doc, '⚠️  Para agregar un producto NUEVO desde la planilla: dejá la columna ID vacía (o en 0), '
             'completá Nombre, precios y stock. En el próximo sync el sistema lo crea automáticamente '
             'en el CRM. Para productos existentes, no toques nunca el ID.',
        color=AMBER)

h2(doc, 'Sync manual')

body(doc, 'Si no querés esperar los 10 minutos automáticos:')

numbered(doc, 'Hacé click en el botón "Sync ahora" en la pantalla del Catálogo.')
numbered(doc, 'El sistema lee la planilla inmediatamente y muestra el resultado.')

h2(doc, 'Exportaciones disponibles')

data_table(doc,
    ['Exportación', 'Formato', 'Cómo usarla'],
    [
        ('Exportar PDF', 'PDF',  'Elegís título y qué productos incluir — descarga listo para enviar a clientes'),
        ('Exportar CSV', 'CSV',  'Descarga todos los productos con nombre, precios, stock, descripción e imagen URL'),
    ],
    col_widths=[4.0, 2.5, 9.0]
)

h2(doc, 'Notificar clientes del catálogo')

body(doc, '"Notificar Clientes" envía automáticamente un email con el link al catálogo público '
         'a todos los leads en estado "cliente". Al terminar muestra cuántos emails fueron encolados.')

h2(doc, 'El Catálogo Público — para compartir con clientes')

body(doc, 'Además del catálogo interno (que ves en el CRM), existe un catálogo público '
         'que cualquier persona puede ver sin necesitar login ni contraseña:')

data_table(doc,
    ['Característica', 'Detalle'],
    [
        ('URL pública',           'kairos.polkorp.com/public/catalog'),
        ('Qué muestra',           'Todos los productos marcados como "Activo" con foto, nombre, descripción y precio minorista'),
        ('Acceso',                'Sin usuario ni contraseña — cualquier persona con el link puede verlo'),
        ('Cómo compartirlo',      'Copiá el link y pegalo en WhatsApp, email o Instagram directamente'),
        ('Actualización',         'Se actualiza en tiempo real — si cambiás un precio en el CRM, el cliente lo ve al instante'),
    ],
    col_widths=[4.0, 11.0]
)

tip_box(doc, '💡  Podés enviar el link del catálogo público a todos tus clientes de una sola vez '
             'usando el botón "Notificar Clientes" o seleccionando leads y usando la acción "Catálogo" '
             'en la barra flotante.',
        color=GOLD)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════
#  SECTION 8 — SCRAPER
# ══════════════════════════════════════════════════════════════════════════

section_header(doc, '8', 'SCRAPER DE LEADS',
               'Búsqueda automática de prospectos y enriquecimiento de contactos')

body(doc, 'El Scraper es la herramienta que llena la base de leads automáticamente. '
         'Busca negocios en Google Places (la misma base de datos que usa Google Maps) '
         'usando búsquedas como "tienda holística Buenos Aires" o "sahumerios Córdoba", '
         'y guarda todos los resultados como leads nuevos en el CRM.')

body(doc, 'El Enriquecedor es un segundo proceso que toma los leads ya cargados y visita '
         'sus sitios web para encontrar datos de contacto que no estaban disponibles en Google Places: '
         'emails, teléfonos adicionales, Instagram, WhatsApp.')

h2(doc, 'Lanzar el Scraper de Google Places')

numbered(doc, 'Hacé click en "Iniciar Scraper".')
numbered(doc, 'El sistema lanza 20 búsquedas predefinidas: "sahumerios Argentina", '
              '"tienda holística Buenos Aires", "santería Córdoba", etc.')
numbered(doc, 'Una barra de progreso muestra el avance en tiempo real (query por query).')
numbered(doc, 'Al finalizar se muestra cuántos leads nuevos fueron encontrados y agregados.')

tip_box(doc, '⚠️  Solo puede correr un job a la vez. Si hay uno activo, el sistema bloquea '
             'el inicio de otro hasta que termine o se cancele.',
        color=AMBER)

h2(doc, 'Enriquecer contactos — mejorar los datos de email y redes')

body(doc, '"Enriquecer con Email" es el proceso que visita automáticamente el sitio web de cada lead '
         'y extrae información de contacto que no tenía. '
         'Ideal para ejecutar después de un scraping cuando muchos leads no tienen email.')

numbered(doc, 'Asegurate de haber corrido el Scraper primero para tener leads con sitio web.')
numbered(doc, 'Hacé click en "Enriquecer con Email".')
numbered(doc, 'El sistema procesa todos los leads que tienen website pero no tienen email.')
numbered(doc, 'Para cada sitio visita las páginas principales buscando datos de contacto.')
numbered(doc, 'Al terminar muestra cuántos leads recibieron datos nuevos.')

data_table(doc,
    ['Dato extraído', 'De dónde lo saca'],
    [
        ('Email',     'mailto: links, JSON-LD, metadatos, footer, texto de la página'),
        ('Instagram', 'Links a instagram.com en la página'),
        ('WhatsApp',  'Links wa.me o whatsapp.com'),
        ('Teléfono',  'Links tel:, datos estructurados, footer'),
    ],
    col_widths=[4.0, 11.0]
)

tip_box(doc, '💡  Después de enriquecer, los leads "Sin Email" del Dashboard van a bajar. '
             'Revisá el Score IA de los leads — los que pasaron de rojo a verde son los que '
             'más datos nuevos recibieron y son prioridad para la próxima campaña.',
        color=GOLD)

h2(doc, 'Flujo completo: del scraping a la campaña')

flow_table(doc, [
    ('🔍', 'Scraper'),
    ('✉', 'Enriquecer'),
    ('⭐', 'Filtrar'),
    ('📢', 'Campaña'),
    ('📱', 'Seguim. WA'),
])

data_table(doc,
    ['Paso', 'Qué hacer'],
    [
        ('1 — Scraper',      'Lanzar el Scraper para agregar negocios nuevos desde Google Places'),
        ('2 — Enriquecer',   'Correr "Enriquecer con Email" para extraer emails de los sitios web'),
        ('3 — Filtrar',      'En Leads, activar "Solo con email", revisar scores y elegir el segmento'),
        ('4 — Campaña',      'Crear campaña con los filtros del segmento y enviar el email'),
        ('5 — Seguim. WA',   'Usar "Seguimiento WA" para contactar los que abrieron pero no respondieron'),
    ],
    col_widths=[4.0, 11.0]
)

h2(doc, 'Historial de jobs')

body(doc, 'La tabla inferior muestra todos los trabajos anteriores con:')
bullet(doc, 'Fecha y hora de inicio y fin.')
bullet(doc, 'Estado: completado / error / corriendo / pendiente.')
bullet(doc, 'Total encontrados y nuevos agregados.')
bullet(doc, 'Mensaje de error si el job falló.')
bullet(doc, 'Botón "Cancelar" para detener un job en curso.')

tip_box(doc, '💡  Los jobs tienen un timeout automático de 30 minutos. Si un job se cuelga '
             'sin actividad por más de ese tiempo, el sistema lo marca como fallido '
             'automáticamente.',
        color=GOLD)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════
#  FLUJO DE TRABAJO RECOMENDADO
# ══════════════════════════════════════════════════════════════════════════

section_header(doc, '✦', 'FLUJO DE TRABAJO RECOMENDADO',
               'Del prospecto al cliente — proceso completo')

body(doc, 'Este es el ciclo típico de trabajo desde la búsqueda de nuevos leads '
         'hasta cerrar una venta:')

doc.add_paragraph()

flow_table(doc, [
    ('🔍', 'Scraper'),
    ('✉', 'Enriquecer'),
    ('⭐', 'Priorizar'),
    ('📢', 'Campaña'),
    ('📱', 'Seguim. WA'),
    ('📦', 'Orden'),
    ('📋', 'Catálogo'),
])

data_table(doc,
    ['Paso', 'Acción', 'Sección'],
    [
        ('1', 'Lanzar el Scraper para buscar nuevos leads en Google Places',              'Scraper'),
        ('2', 'Correr el Enriquecedor para extraer emails de los sitios web',             'Scraper'),
        ('3', 'Revisar la lista de Leads, filtrar por score y priorizar los mejores',     'Leads'),
        ('4', 'Enviar campaña de email o catálogo al segmento elegido',                   'Campañas'),
        ('5', 'Usar "Seguimiento WA" para contactar los que no respondieron',             'Campañas'),
        ('6', 'Mover los leads interesados en el Pipeline hasta llegar a "Cliente"',      'Pipeline'),
        ('7', 'Crear la orden cuando el cliente confirma el pedido',                      'Órdenes'),
        ('8', 'Mantener precios y stock actualizados en el Catálogo vía Google Sheets',   'Catálogo'),
    ],
    col_widths=[1.2, 10.3, 3.5]
)

# ── FAQ ────────────────────────────────────────────────────────────────────

h2(doc, 'Preguntas Frecuentes')

faqs = [
    ('¿Puedo usar el sistema desde el celular?',
     'Sí. El diseño es responsivo. En pantallas chicas la tabla de leads se convierte en tarjetas '
     'apiladas y el Pipeline/Órdenes scrollean horizontalmente.'),
    ('¿Cómo actualizo los precios de muchos productos a la vez?',
     'Conectá una planilla de Google Sheets desde el Catálogo. Una vez conectada, editá los precios '
     'directamente en la planilla y el sistema los aplica automáticamente cada 10 minutos.'),
    ('¿Cómo creo una campaña de email para un segmento específico?',
     'En Campañas → Nueva Campaña. En el Paso 1 aplicá los filtros (Provincia, Rubro, Estado, '
     'Solo con email). En el Paso 2 escribí o generá el texto con IA. En el Paso 3 revisá '
     'la cantidad de destinatarios y enviá.'),
    ('¿El scraper agrega duplicados?',
     'No. Antes de insertar un lead, verifica que no exista ya una empresa con el mismo nombre y '
     'tipo de cliente.'),
    ('¿Cómo sé si un email llegó y fue abierto?',
     'Desde el detalle de una campaña (click en el nombre) podés ver enviados, abiertos, clicks '
     'y respondidos por lead. Requiere que Brevo esté configurado correctamente.'),
    ('¿Cómo le mando el catálogo a un cliente específico?',
     'Desde Leads, seleccioná ese lead, hacé click en "Catálogo" en la barra flotante. '
     'También podés exportar el PDF del Catálogo y adjuntarlo manualmente por WhatsApp.'),
    ('¿Cómo mejoro los datos de mis contactos?',
     'Usá el Enriquecedor en la sección Scraper → "Enriquecer con Email". '
     'El sistema visita los sitios web de los leads y extrae emails, teléfonos e Instagram automáticamente.'),
    ('¿Cómo armo una lista con precios para cotizar?',
     'Filtrá los leads en la sección Leads (por provincia, rubro, estado), activá "Solo con email" '
     'y exportá el CSV. Luego exportá también el CSV del Catálogo. Combiná ambos en Excel '
     'o enviá el link del catálogo público directo por WhatsApp.'),
    ('¿Qué pasa si el scraper se cuelga?',
     'Tiene timeout automático de 30 minutos. También podés cancelarlo manualmente desde el historial.'),
    ('¿Cómo registro que hablé con un lead?',
     'Abrí el detalle del lead (click en su nombre), bajá a "Timeline de notas" y agregá una nota '
     'con el resumen de la conversación. Así queda registrado con fecha y hora.'),
]

for q, a in faqs:
    # Question
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(8)
    pq.paragraph_format.space_after  = Pt(2)
    rq = pq.add_run('❓  ' + q)
    rq.bold            = True
    rq.font.size       = Pt(10)
    rq.font.color.rgb  = BROWN_DARK
    # Answer
    pa = doc.add_paragraph()
    pa.paragraph_format.space_before = Pt(0)
    pa.paragraph_format.space_after  = Pt(4)
    set_left_indent(pa, 0.8)
    ra = pa.add_run(a)
    ra.font.size = Pt(9.5)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
#  GLOSARIO
# ══════════════════════════════════════════════════════════════════════════

section_header(doc, '✦', 'GLOSARIO',
               'Definición de todos los términos técnicos y de negocio del sistema')

body(doc, 'Si en algún momento no entendés un término al usar el sistema, buscalo acá:')

data_table(doc,
    ['Término', 'Definición'],
    [
        ('Brevo (ex Sendinblue)',
         'Servicio externo de envío de emails masivos. El CRM lo usa para enviar campañas. '
         'Sin Brevo configurado, las campañas de email no llegan.'),
        ('CRM',
         'Customer Relationship Management — sistema de gestión de relaciones con clientes. '
         'Permite registrar y seguir todos los contactos y oportunidades de venta.'),
        ('CSV',
         'Comma-Separated Values — formato de planilla de texto plano compatible con Excel y Google Sheets. '
         'Se usa para importar/exportar leads y productos en masa.'),
        ('Dashboard',
         'Pantalla de inicio del sistema con un resumen visual del estado del negocio en tiempo real.'),
        ('Email de campaña',
         'Email enviado masivamente a un grupo de leads. Diferente a un email individual — '
         'se personaliza con variables como {nombre} y {empresa}.'),
        ('Embudo de ventas',
         'El proceso por etapas desde que un contacto es Nuevo hasta que se convierte en Cliente. '
         'El Pipeline muestra el embudo en formato visual Kanban.'),
        ('Enriquecimiento',
         'Proceso automático que visita el sitio web de cada lead y extrae emails, teléfonos, '
         'Instagram y WhatsApp que no estaban registrados.'),
        ('Google Places',
         'Base de datos de negocios de Google (la misma que usa Google Maps). '
         'El Scraper busca prospectos nuevos en esta base.'),
        ('Google Sheets',
         'Planilla de cálculo en la nube de Google (similar a Excel online). '
         'El catálogo puede sincronizarse automáticamente desde una planilla de Sheets.'),
        ('Job (trabajo)',
         'Una tarea en ejecución en segundo plano: el Scraper corre como un "job". '
         'Se puede ver su progreso y cancelarlo desde el historial.'),
        ('Kanban',
         'Método visual de gestión con columnas y tarjetas. El Pipeline y las Órdenes usan tableros Kanban.'),
        ('Lead',
         'Un contacto o prospecto — negocio que podría comprar productos Kairos. '
         'Cada lead tiene empresa, teléfono, email, ciudad, provincia y rubro.'),
        ('Mayorista',
         'Distribuidor o proveedor que compra en volumen. Se gestiona igual que los leads '
         'pero en su propia sección con listas de precios mayoristas.'),
        ('Pipeline',
         'Vista del embudo de ventas en formato Kanban. Muestra los leads organizados '
         'en columnas según su estado comercial.'),
        ('Rubro',
         'Tipo de negocio del lead: tienda holística, sahumerios, santería, etc. '
         'Se usa como filtro para segmentar campañas.'),
        ('Score IA',
         'Puntaje del 0 al 10 que el sistema calcula automáticamente según cuántos datos '
         'tiene el lead: verde (7-10), amarillo (4-6), rojo (0-3).'),
        ('Scraper',
         'Programa que busca negocios automáticamente en Google Places y los agrega a la base de leads.'),
        ('Segmento',
         'Grupo de leads filtrados por criterios (provincia, rubro, estado, tiene email). '
         'Las campañas se envían a un segmento.'),
        ('Sync / Sincronizar',
         'Proceso de actualización automática entre el CRM y un sistema externo (Google Sheets). '
         'Los cambios en la planilla se aplican al CRM en el próximo ciclo de sync.'),
        ('Tasa de apertura',
         'Porcentaje de destinatarios que abrieron el email. Una tasa normal es 20-40%.'),
        ('Toggle',
         'Interruptor de on/off en la interfaz — un botón que se activa o desactiva con un click.'),
        ('{empresa} / {nombre}',
         'Variables de personalización en el texto de emails. El sistema las reemplaza '
         'automáticamente con el nombre real de cada destinatario.'),
    ],
    col_widths=[4.0, 11.0]
)

# ── FOOTER NOTE ────────────────────────────────────────────────────────────
doc.add_paragraph()
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.paragraph_format.space_before = Pt(20)
rf = pf.add_run('kairos.polkorp.com  ·  Kairos Distribuidora  ·  Guía del Sistema CRM')
rf.font.size      = Pt(8.5)
rf.font.color.rgb = BROWN_MID
rf.italic         = True

# ══════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════
out = '/home/user/kairos/GUIA_KAIROS_CRM.docx'
doc.save(out)
print(f'Saved → {out}')
